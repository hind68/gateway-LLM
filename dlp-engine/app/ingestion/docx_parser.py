from docx import Document

from app.ingestion.ooxml_media import extract_text_from_embedded_images


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    # Headers/footers aren't covered by doc.paragraphs above - letterheads
    # and "Confidential" footers commonly carry contact info or markings
    # that would otherwise be silently skipped.
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                if p.text.strip():
                    parts.append(p.text)

    # A pasted screenshot or scanned photo (e.g. of an ID card) inside the
    # document has text too, same reasoning as pdf_parser.py's OCR pass.
    parts.extend(extract_text_from_embedded_images(path, "word/media/"))

    return "\n".join(parts)
