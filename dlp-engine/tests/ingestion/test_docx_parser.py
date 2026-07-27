import os
import tempfile

from docx import Document
from PIL import Image, ImageDraw, ImageFont

from app.ingestion.docx_parser import extract_text_from_docx


def _render_text_image(text: str, path: str):
    img = Image.new("RGB", (600, 100), color="white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 32)
    except Exception:
        font = ImageFont.load_default()
    draw.text((10, 30), text, fill="black", font=font)
    img.save(path)


def test_extracts_paragraph_and_table_text():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc = Document()
    doc.add_paragraph("Confidential memo")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "john@company.com"
    doc.save(path)

    text = extract_text_from_docx(path)
    os.remove(path)

    assert "Confidential memo" in text
    assert "john@company.com" in text

def test_extracts_header_and_footer_text():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Acme Corp Letterhead"
    doc.sections[0].footer.paragraphs[0].text = "Confidential - internal use only"
    doc.add_paragraph("Body text")
    doc.save(path)

    text = extract_text_from_docx(path)
    os.remove(path)

    assert "Acme Corp Letterhead" in text
    assert "Confidential - internal use only" in text

def test_ocrs_embedded_image():
    img_path = tempfile.mktemp(suffix=".png")
    _render_text_image("CONFIDENTIAL EMAIL", img_path)

    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc = Document()
    doc.add_paragraph("See attached scan")
    doc.add_picture(img_path)
    doc.save(docx_path)

    text = extract_text_from_docx(docx_path)
    os.remove(docx_path)
    os.remove(img_path)

    assert "See attached scan" in text
    assert "CONFIDENTIAL" in text.upper()

def test_docx_with_no_images_still_works():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc = Document()
    doc.add_paragraph("Just plain text, no images")
    doc.save(path)

    text = extract_text_from_docx(path)
    os.remove(path)

    assert "Just plain text, no images" in text
